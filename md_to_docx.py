"""
md_to_docx.py - Convert Markdown files to Word documents (.docx)

Handles: headings, bold, italic, bold+italic, inline code, code blocks,
bullet lists, numbered lists, nested lists, tables, blockquotes, links,
horizontal rules, and plain paragraphs.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from bs4 import BeautifulSoup, Comment, NavigableString, Tag


def _add_inline(para, element):
    """
    Recursively walk an HTML element's children and add runs to the
    paragraph with the correct Bold / Italic / Code formatting.
    """
    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                para.add_run(text)

        elif isinstance(child, Tag):
            name = child.name

            if name in ('strong', 'b'):
                # May contain nested <em>
                for sub in child.children:
                    if isinstance(sub, NavigableString):
                        run = para.add_run(str(sub))
                        run.bold = True
                    elif isinstance(sub, Tag) and sub.name in ('em', 'i'):
                        run = para.add_run(sub.get_text())
                        run.bold = True
                        run.italic = True
                    else:
                        run = para.add_run(sub.get_text())
                        run.bold = True

            elif name in ('em', 'i'):
                for sub in child.children:
                    if isinstance(sub, NavigableString):
                        run = para.add_run(str(sub))
                        run.italic = True
                    elif isinstance(sub, Tag) and sub.name in ('strong', 'b'):
                        run = para.add_run(sub.get_text())
                        run.bold = True
                        run.italic = True
                    else:
                        run = para.add_run(sub.get_text())
                        run.italic = True

            elif name == 'code':
                run = para.add_run(child.get_text())
                run.font.name = 'Courier New'
                run.font.size = Pt(10)

            elif name == 'a':
                # Show link text, underlined
                run = para.add_run(child.get_text())
                run.underline = True

            elif name == 'br':
                para.add_run('\n')

            else:
                # Unknown inline tag — just recurse
                _add_inline(para, child)


def _process_list(doc, list_el, level=0):
    """Handle <ul> and <ol>, including nested lists."""
    ordered = list_el.name == 'ol'
    style = 'List Number' if ordered else 'List Bullet'

    for child in list_el.children:
        if not isinstance(child, Tag) or child.name != 'li':
            continue

        # Separate nested lists from the li's own text content
        nested_lists = child.find_all(['ul', 'ol'], recursive=False)

        # Build a copy of li without nested lists for text extraction
        li_clone = BeautifulSoup(str(child), 'html.parser').find('li')
        for nl in li_clone.find_all(['ul', 'ol']):
            nl.decompose()

        para = doc.add_paragraph(style=style)
        para.paragraph_format.left_indent = Inches(0.25 * level)
        _add_inline(para, li_clone)

        for nested in nested_lists:
            _process_list(doc, nested, level + 1)


def _process_table(doc, table_el):
    """Add a Markdown table as a proper Word table with a border grid."""
    rows = table_el.find_all('tr')
    if not rows:
        return

    max_cols = max(len(r.find_all(['td', 'th'])) for r in rows)
    if max_cols == 0:
        return

    tbl = doc.add_table(rows=0, cols=max_cols)
    tbl.style = 'Table Grid'

    for row_el in rows:
        cells_el = row_el.find_all(['td', 'th'])
        word_row = tbl.add_row()
        for i, cell_el in enumerate(cells_el):
            if i >= max_cols:
                break
            cell = word_row.cells[i]
            # Clear default empty paragraph
            para = cell.paragraphs[0]
            para.clear()
            _add_inline(para, cell_el)
            if cell_el.name == 'th':
                for run in para.runs:
                    run.bold = True


def convert_md_to_docx(md_path, docx_path=None):
    """
    Convert a Markdown file to a Word (.docx) document.
    The .docx is saved next to the original unless docx_path is given.
    Returns the Path of the saved file.
    """
    import markdown as md_lib

    md_path = Path(md_path)
    text = md_path.read_text(encoding='utf-8', errors='replace')

    # Strip YAML / TOML frontmatter (--- ... --- blocks at the top)
    text = re.sub(r'^---\s*\n.*?\n---\s*\n', '', text, flags=re.DOTALL)

    # Convert Markdown → HTML
    html = md_lib.markdown(
        text,
        extensions=['tables', 'fenced_code', 'nl2br', 'sane_lists']
    )

    soup = BeautifulSoup(html, 'html.parser')
    doc = Document()

    # Readable margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    for element in soup.children:
        # Skip HTML comments (e.g. <!-- Slide number: 1 --> from PPTX conversions)
        if isinstance(element, Comment):
            continue
        if isinstance(element, NavigableString):
            stripped = str(element).strip()
            if stripped:
                doc.add_paragraph(stripped, style='Normal')
            continue
        if not isinstance(element, Tag):
            continue

        tag = element.name

        # ── Headings ─────────────────────────────────────────────────────────
        if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            level = int(tag[1])
            para = doc.add_paragraph(style=f'Heading {min(level, 9)}')
            _add_inline(para, element)

        # ── Normal paragraph ──────────────────────────────────────────────────
        elif tag == 'p':
            para = doc.add_paragraph(style='Normal')
            _add_inline(para, element)

        # ── Lists ─────────────────────────────────────────────────────────────
        elif tag in ('ul', 'ol'):
            _process_list(doc, element, level=0)

        # ── Tables ───────────────────────────────────────────────────────────
        elif tag == 'table':
            _process_table(doc, element)
            doc.add_paragraph()   # breathing room after table

        # ── Fenced / indented code blocks ─────────────────────────────────────
        elif tag == 'pre':
            code_el = element.find('code')
            code_text = code_el.get_text() if code_el else element.get_text()
            para = doc.add_paragraph(style='Normal')
            para.paragraph_format.left_indent = Inches(0.5)
            run = para.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

        # ── Block quotes ──────────────────────────────────────────────────────
        elif tag == 'blockquote':
            para = doc.add_paragraph(style='Normal')
            para.paragraph_format.left_indent = Inches(0.5)
            run = para.add_run('| ')
            run.italic = True
            _add_inline(para, element)

        # ── Horizontal rule ───────────────────────────────────────────────────
        elif tag == 'hr':
            para = doc.add_paragraph('─' * 60, style='Normal')

        # ── Fallback ──────────────────────────────────────────────────────────
        else:
            text_content = element.get_text().strip()
            if text_content:
                doc.add_paragraph(text_content, style='Normal')

    if docx_path is None:
        docx_path = md_path.with_suffix('.docx')

    doc.save(str(docx_path))
    return Path(docx_path)
